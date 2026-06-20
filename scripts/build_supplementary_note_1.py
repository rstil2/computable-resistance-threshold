#!/usr/bin/env python3
"""Build Supplementary Note 1 as DOCX and PDF."""

from __future__ import annotations

import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "manuscript"))

from supplementary_note_1_content import (  # noqa: E402
    EMPIRICAL_PAPER,
    INTRO,
    SECTIONS,
    SUBTITLE,
    TITLE,
)
from content import AUTHORS, AFFILIATION  # noqa: E402

OUT_DOCX = ROOT / "manuscript" / "SUPPLEMENTARY_NOTE_1_PAC.docx"
OUT_PDF = ROOT / "manuscript" / "SUPPLEMENTARY_NOTE_1_PAC.pdf"


def setup_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, indent: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.left_indent = Pt(18)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def build_docx() -> Path:
    doc = setup_doc()

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(TITLE)
    tr.bold = True
    tr.font.size = Pt(14)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sp.add_run(SUBTITLE)
    sr.italic = True
    sr.font.size = Pt(12)

    ep = doc.add_paragraph()
    ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    er = ep.add_run(f"Companion to: {EMPIRICAL_PAPER}")
    er.font.size = Pt(10)
    er.italic = True

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ap.add_run(f"{AUTHORS} · {AFFILIATION}")
    ar.font.size = Pt(10)

    doc.add_paragraph()

    for para in INTRO:
        add_para(doc, para, italic=para.startswith("Stillwell,"))

    for heading, paragraphs in SECTIONS:
        add_heading(doc, heading, level=2)
        for para in paragraphs:
            indent = para.startswith("    ") or para.endswith(",")
            text = para.strip()
            add_para(doc, text, indent=indent)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    return OUT_DOCX


def build_pdf(docx_path: Path) -> Path:
    soffice = shutil.which("soffice")
    if not soffice:
        raise RuntimeError("LibreOffice (soffice) not found; cannot export PDF")

    if OUT_PDF.exists():
        OUT_PDF.unlink()

    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf", "--outdir", str(OUT_DOCX.parent), str(docx_path)],
        check=True,
        capture_output=True,
        text=True,
    )
    if not OUT_PDF.exists():
        raise RuntimeError(f"PDF conversion failed: expected {OUT_PDF}")
    return OUT_PDF


def build() -> tuple[Path, Path]:
    docx = build_docx()
    pdf = build_pdf(docx)
    return docx, pdf


if __name__ == "__main__":
    docx, pdf = build()
    print(f"Wrote {docx}")
    print(f"Wrote {pdf}")
