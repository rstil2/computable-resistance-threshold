#!/usr/bin/env python3
"""Build complete Nature Cancer / Nat Comm manuscript DOCX with embedded figures."""

from __future__ import annotations

import sys
from pathlib import Path

import fitz  # pymupdf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(ROOT / "manuscript"))

from content import (  # noqa: E402
    ABSTRACT,
    ACKNOWLEDGEMENTS,
    AFFILIATION,
    AUTHOR_CONTRIBUTIONS,
    AUTHORS,
    COMPETING_INTERESTS,
    CORRESPONDING,
    DATA_AVAILABILITY,
    DISCUSSION,
    ETHICS,
    EXTENDED_DATA_TABLES,
    GITHUB_REPO,
    INTRODUCTION,
    KEY_POINTS,
    KEYWORDS,
    METHODS,
    REFERENCES,
    RESULTS,
    TITLE,
)

OUT = ROOT / "manuscript" / "MANUSCRIPT_draft.docx"

FIGURES = {
    "Figure 1": ROOT / "data/fig1/fig1_clinical_report.png",
    "Figure 2": ROOT / "data/fig2_cross_cancer/fig2_cross_cancer.png",
    "Figure 3": ROOT / "data/fig3_targeted/fig3_targeted_therapy.png",
    "Figure 4": ROOT / "data/fig4_utility/fig4_clinical_utility.png",
    "Extended Data Fig. 1": ROOT / "data/ed_tcga/ed_tcga_pfs_km.png",
}

FIG_WIDTH = Inches(6.75)


def resolve_image(path: Path, cache_dir: Path) -> Path:
    if path.suffix.lower() == ".png" and path.exists():
        return path
    if path.suffix.lower() == ".pdf" and path.exists():
        cache_dir.mkdir(parents=True, exist_ok=True)
        out = cache_dir / (path.stem + ".png")
        if not out.exists() or out.stat().st_mtime < path.stat().st_mtime:
            doc = fitz.open(path)
            pix = doc[0].get_pixmap(matrix=fitz.Matrix(4, 4), alpha=False)
            pix.save(str(out))
            doc.close()
        return out
    png = path.with_suffix(".png")
    if png.exists():
        return png
    raise FileNotFoundError(path)


def setup_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)
    return doc


def add_para(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 11) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Times New Roman"


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        for run in p.runs:
            run.font.size = Pt(11)


def add_figure(doc: Document, label: str, image_path: Path, legend: str) -> None:
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=FIG_WIDTH)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r1 = cap.add_run(f"{label}. ")
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = cap.add_run(legend.replace("**", ""))
    r2.font.size = Pt(10)
    r2.italic = True
    cap.paragraph_format.space_after = Pt(14)


def render_blocks(doc: Document, blocks: list, figure_paths: dict[str, Path]) -> None:
    for block in blocks:
        if block[0] == "heading2":
            add_heading(doc, block[1], level=2)
        elif block[0] == "para":
            add_para(doc, block[1])
        elif block[0] == "figure":
            _, label, legend = block
            add_figure(doc, label, figure_paths[label], legend)


def build() -> Path:
    cache = ROOT / "manuscript" / "_figure_cache"
    figure_paths = {k: resolve_image(v, cache) for k, v in FIGURES.items()}

    doc = setup_doc()

    # Title block
    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(TITLE)
    tr.bold = True
    tr.font.size = Pt(16)

    ap = doc.add_paragraph()
    ap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = ap.add_run(AUTHORS)
    ar.font.size = Pt(11)

    af = doc.add_paragraph()
    af.alignment = WD_ALIGN_PARAGRAPH.CENTER
    afr = af.add_run(AFFILIATION)
    afr.font.size = Pt(10)
    afr.italic = True

    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = cp.add_run(CORRESPONDING)
    cr.font.size = Pt(10)
    cr.italic = True

    doc.add_paragraph()

    # Abstract
    add_heading(doc, "Abstract", 1)
    for para in ABSTRACT:
        add_para(doc, para)
    add_para(doc, f"Keywords: {KEYWORDS}", italic=True, size=10)

    # Key points (Nature Cancer style summary)
    add_heading(doc, "Key points", 1)
    add_bullets(doc, KEY_POINTS)

    # Introduction
    add_heading(doc, "Introduction", 1)
    for para in INTRODUCTION:
        add_para(doc, para)

    # Results
    add_heading(doc, "Results", 1)
    render_blocks(doc, RESULTS, figure_paths)

    # Discussion
    add_heading(doc, "Discussion", 1)
    for para in DISCUSSION:
        add_para(doc, para)

    # Methods
    add_heading(doc, "Methods", 1)
    render_blocks(doc, METHODS, figure_paths)

    # Extended Data tables (text summary)
    add_heading(doc, "Extended Data", 1)
    for table_title, table_sub, table_body in EXTENDED_DATA_TABLES:
        add_heading(doc, table_title, level=2)
        add_para(doc, table_sub, bold=True, italic=True)
        add_para(doc, table_body)

    # Back matter
    add_heading(doc, "Data availability", 1)
    add_para(doc, DATA_AVAILABILITY)

    add_heading(doc, "Code availability", 1)
    add_para(
        doc,
        f"Analysis code, locked configuration (config/), and figure regeneration scripts: "
        f"{GITHUB_REPO}. Pre-registration: https://osf.io/kp5jf."
    )

    add_heading(doc, "Ethics statement", 1)
    add_para(doc, ETHICS)

    add_heading(doc, "Acknowledgements", 1)
    add_para(doc, ACKNOWLEDGEMENTS)

    add_heading(doc, "Author contributions", 1)
    add_para(doc, AUTHOR_CONTRIBUTIONS)

    add_heading(doc, "Competing interests", 1)
    add_para(doc, COMPETING_INTERESTS)

    add_heading(doc, "References", 1)
    for i, ref in enumerate(REFERENCES, 1):
        p = doc.add_paragraph(f"{i}. {ref}")
        p.paragraph_format.space_after = Pt(3)
        for run in p.runs:
            run.font.size = Pt(10)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    try:
        print(f"Wrote {build()}")
    except FileNotFoundError as e:
        print(f"Missing figure: {e}", file=sys.stderr)
        print("Run: python scripts/run_nat_cancer_figures.py", file=sys.stderr)
        sys.exit(1)
